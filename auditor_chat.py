import streamlit as st
import google.generativeai as genai
import PyPDF2
from docx import Document 
import io

# ==========================================
# 1. CONFIGURAÇÕES VISUAIS DA PÁGINA
# ==========================================
st.set_page_config(page_title="Auditor Master AI", page_icon="⚖️", layout="wide")

if "chat_session" not in st.session_state:
    st.session_state.chat_session = None
if "mensagens_chat" not in st.session_state:
    st.session_state.mensagens_chat = []
if "relatorio_denso" not in st.session_state:
    st.session_state.relatorio_denso = ""

# ==========================================
# 2. FUNÇÃO PARA GERAR O ARQUIVO WORD
# ==========================================
def gerar_docx(texto_markdown):
    doc = Document()
    doc.add_heading('Relatório de Auditoria Jurídica AI', 0)
    
    linhas = texto_markdown.split('\n')
    for linha in linhas:
        texto_limpo = linha.replace('**', '').replace('*', '•')
        if linha.startswith('###'):
            doc.add_heading(texto_limpo.replace('###', '').strip(), level=3)
        elif linha.startswith('##'):
            doc.add_heading(texto_limpo.replace('##', '').strip(), level=2)
        elif linha.startswith('#'):
            doc.add_heading(texto_limpo.replace('#', '').strip(), level=1)
        else:
            doc.add_paragraph(texto_limpo)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. BARRA LATERAL 
# ==========================================
with st.sidebar:
    st.header("⚙️ Configurações do Sistema")
    chave_api = st.text_input("Chave da API do Google:", type="password")
    st.divider()
    st.subheader("📂 Documentos")
    arquivos_pdf = st.file_uploader("Anexe os contratos em PDF", type=["pdf"], accept_multiple_files=True)
    btn_processar = st.button("⚖️ Executar Auditoria Completa", type="primary", use_container_width=True)

# ==========================================
# 4. O MOTOR DE INTELIGÊNCIA 
# ==========================================
if btn_processar:
    if not chave_api:
        st.sidebar.error("⚠️ Insira a Chave de API primeiro.")
    elif not arquivos_pdf:
        st.sidebar.warning("⚠️ Anexe pelo menos um contrato.")
    else:
        with st.spinner("Extraindo textos e executando análise profunda..."):
            try:
                texto_completo = ""
                for pdf in arquivos_pdf:
                    leitor = PyPDF2.PdfReader(pdf)
                    for pagina in leitor.pages:
                        texto_completo += pagina.extract_text() + "\n"
                
                genai.configure(api_key=chave_api)
                modelo = genai.GenerativeModel('gemini-2.5-flash')
                
                # CÉREBRO DO ORÁCULO CORRIGIDO E LIMPO!
                prompt_auditoria = f"""
                Você é o Auditor Jurídico PRO, um auditor jurídico de elite. 
                Sua análise deve ser cirúrgica. Mantenha sobriedade máxima.
                Use negrito para destacar valores, prazos e nomes importantes.
                
                Estruture seu relatório OBRIGATORIAMENTE com:
                1. RESUMO EXECUTIVO do Acordo.
                2. PONTOS CRÍTICOS (Os principais riscos jurídicos ou financeiros).
                3. SUGESTÃO ESTRATÉGICA (Como consertar ou blindar as cláusulas).
                
                CONTRATO:
                {texto_completo}
                """
                
                resposta_relatorio = modelo.generate_content(prompt_auditoria)
                st.session_state.relatorio_denso = resposta_relatorio.text
                
                st.session_state.chat_session = modelo.start_chat(history=[
                    {"role": "user", "parts": [f"Aqui está o contrato integral:\n\n{texto_completo}\n\nResponda apenas a dúvidas específicas sobre o documento."]},
                    {"role": "model", "parts": ["Entendido. Estou à disposição para responder dúvidas pontuais ou explicar termos deste contrato."]}
                ])
                
                st.session_state.mensagens_chat = [
                    {"role": "assistant", "content": "Olá! Viu algo estranho no relatório? Me pergunte qual cláusula alterar ou tire suas dúvidas abaixo."}
                ]
                
            except Exception as e:
                st.error(f"Erro no processamento: {e}")

# ==========================================
# 5. A INTERFACE PRINCIPAL (UX REVISADA)
# ==========================================
st.title("⚖️ Auditor Jurídico Master AI")

if st.session_state.relatorio_denso:
    
    st.success("✅ **Auditoria concluída com sucesso!** Navegue pelas DUAS ABAS abaixo para ver o Relatório ou Conversar com a IA.")
    
    aba_relatorio, aba_chat = st.tabs(["📑 1. O RELATÓRIO COMPLETO", "💬 2. CHAT TIRA-DÚVIDAS"])
    
    with aba_relatorio:
        st.markdown(st.session_state.relatorio_denso)
        st.divider()
        
        arquivo_word = gerar_docx(st.session_state.relatorio_denso)
        st.download_button(
            label="📄 Baixar Relatório Completo (Formato Word / .docx)",
            data=arquivo_word,
            file_name="Relatorio_Auditoria.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
    with aba_chat: # <--- OLHA ELE AQUI!
        st.caption("Faça perguntas livres sobre os documentos processados.")
        
        # --- O NOVO BOTÃO DE EXPORTAR CHAT ---
        if st.session_state.mensagens_chat:
            texto_chat = "--- HISTÓRICO DA AUDITORIA (CHAT) ---\n\n"
            for msg in st.session_state.mensagens_chat:
                autor = "🧑‍⚖️ Advogado" if msg["role"] == "user" else "🤖 Auditor AI"
                texto_chat += f"{autor}:\n{msg['content']}\n\n"
                
            st.download_button(
                label="📥 Baixar Histórico do Chat (.txt)",
                data=texto_chat,
                file_name="Historico_Chat_Auditoria.txt",
                mime="text/plain"
            )
        st.divider()
        # ------------------------------------
        
        for msg in st.session_state.mensagens_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if prompt := st.chat_input("Ex: Verifique se o nome dos locatários está correto."):
            st.session_state.mensagens_chat.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
                
            with st.chat_message("assistant"):
                with st.spinner("Pensando..."):
                    resposta_chat = st.session_state.chat_session.send_message(prompt)
                    st.markdown(resposta_chat.text)
            st.session_state.mensagens_chat.append({"role": "assistant", "content": resposta_chat.text})

else:
    st.info("👈 Faça o upload dos contratos em PDF na barra lateral e clique em **Executar Auditoria Completa** para começar.")